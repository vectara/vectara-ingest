import logging
logger = logging.getLogger(__name__)
from typing import List, Tuple, Iterator, Dict, Any, Optional
import time
import pandas as pd
import os
from io import StringIO, BytesIO
import gc
import base64
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
import requests
from urllib.parse import urlparse
from dataclasses import dataclass

import pathlib
from pdf2image import convert_from_bytes
from slugify import slugify

from core.summary import TableSummarizer, ImageSummarizer
from core.utils import detect_file_type, markdown_to_df, get_headers, MIN_IMAGE_DIMENSION, release_memory
from core.context_utils import extract_image_context

import unstructured as us
from unstructured.partition.pdf import partition_pdf
from unstructured.partition.html import partition_html
from unstructured.partition.pptx import partition_pptx
from unstructured.partition.docx import partition_docx

from omegaconf import OmegaConf

import nest_asyncio
from llama_parse import LlamaParse

from gmft.pdf_bindings import PyPDFium2Document
from gmft.auto import TableDetector, AutoTableFormatter, AutoFormatConfig

import nltk
nltk.download('punkt_tab', quiet=True)
nltk.download('averaged_perceptron_tagger_eng', quiet=True)

from PIL import ImageFile
ImageFile.LOAD_TRUNCATED_IMAGES = True

MAX_VERBOSE_LENGTH = 1000

import warnings
warnings.filterwarnings(
    "ignore",
    message=".*pin_memory.*no accelerator is found.*"
)


def extract_document_title(filename: str) -> str:
    """
    Extract title from document metadata using appropriate libraries.
    
    Supports:
    - PDF files: Uses pypdf to extract title from PDF metadata
    - DOCX files: Uses python-docx to extract title from document properties  
    - PPTX files: Uses python-pptx to extract title from presentation properties
    - HTML files: Parses the <title> tag from the HTML document
    - Other files: Returns empty (will fallback to filename or Title elements)
    
    Args:
        filename (str): Path to the document file
        
    Returns:
        str: Document title from metadata, or empty string if not found or on error
    """
    if filename.endswith('.pdf'):
        return _extract_pdf_title(filename)
    elif filename.endswith('.docx'):
        return _extract_docx_title(filename)
    elif filename.endswith('.pptx'):
        return _extract_pptx_title(filename)
    elif filename.endswith(('.html', '.htm')):
        return _extract_html_title(filename)
    else:
        # Other files return empty for fallback to Title elements or filename
        return ''


def _extract_pdf_title(filename: str) -> str:
    """Extract title from PDF metadata using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(filename)
        if reader.metadata and reader.metadata.title:
            title = reader.metadata.title.strip()
            if title:  # Only return non-empty titles
                logger.info(f"Extracted PDF metadata title: '{title}' from file {filename}")
                return title
    except Exception as e:
        logger.warning(f"Failed to extract PDF metadata title from {filename}: {e}")
    return ''


def _extract_docx_title(filename: str) -> str:
    """Extract title from DOCX document properties using python-docx."""
    try:
        from docx import Document
        doc = Document(filename)
        if doc.core_properties.title:
            title = doc.core_properties.title.strip()
            if title:  # Only return non-empty titles
                logger.info(f"Extracted DOCX document title: '{title}' from file {filename}")
                return title
    except ImportError:
        logger.debug(f"python-docx not available, skipping DOCX title extraction for {filename}")
    except Exception as e:
        logger.warning(f"Failed to extract DOCX document title from {filename}: {e}")
    return ''


def _extract_pptx_title(filename: str) -> str:
    """Extract title from PPTX presentation properties using python-pptx."""
    try:
        from pptx import Presentation
        prs = Presentation(filename)
        if prs.core_properties.title:
            title = prs.core_properties.title.strip()
            if title:  # Only return non-empty titles
                logger.info(f"Extracted PPTX presentation title: '{title}' from file {filename}")
                return title
    except ImportError:
        logger.debug(f"python-pptx not available, skipping PPTX title extraction for {filename}")
    except Exception as e:
        logger.warning(f"Failed to extract PPTX presentation title from {filename}: {e}")
    return ''


def _extract_html_title(filename: str) -> str:
    """Extract title from HTML <title> tag."""
    from html.parser import HTMLParser

    class _TitleParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self._in_title = False
            self.title = ''

        def handle_starttag(self, tag, attrs):
            if tag.lower() == 'title':
                self._in_title = True

        def handle_endtag(self, tag):
            if tag.lower() == 'title':
                self._in_title = False

        def handle_data(self, data):
            if self._in_title:
                self.title += data

    try:
        with open(filename, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(8192)  # <title> is always in the <head>, first 8KB is enough
        parser = _TitleParser()
        parser.feed(content)
        title = parser.title.strip()
        if title:
            logger.info(f"Extracted HTML title: '{title}' from file {filename}")
            return title
    except Exception as e:
        logger.warning(f"Failed to extract HTML title from {filename}: {e}")
    return ''


@dataclass
class ParsedDocument:
    """
    Unified document representation with content in document order
    """
    title: str
    content_stream: List[Tuple[Any, Dict[str, Any]]]  # All content elements in document order
    tables: List[Tuple]  # Detailed table data for structured indexing
    image_bytes: List[Tuple[str, bytes]]  # Image binary data as (image_id, binary_data)
    
    def get_texts(self) -> List[Tuple[str, Dict[str, Any]]]:
        """Get only text elements from content stream"""
        return [(content, metadata) for content, metadata in self.content_stream 
                if metadata.get('element_type') == 'text']
    
    def get_images(self) -> List[Tuple[str, Dict[str, Any]]]:
        """Get only image elements from content stream"""
        return [(content, metadata) for content, metadata in self.content_stream 
                if metadata.get('element_type') == 'image']
    
    def get_all_content(self) -> List[Tuple[Any, Dict[str, Any]]]:
        """Get all content elements in document order"""
        return self.content_stream.copy()
    
    def to_legacy_format(self) -> Tuple[str, List[Tuple], List[Tuple], List[Tuple]]:
        """Convert to legacy (title, texts, tables, images) format for backward compatibility"""
        return (self.title, self.get_texts(), self.tables, self.get_images())
    
    def __iter__(self):
        """Allow direct unpacking: title, texts, tables, images = parsed_doc"""
        return iter(self.to_legacy_format())


class DocumentParser():
    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = {},
        parse_tables: bool = False,
        enable_gmft: bool = False,
        do_ocr: bool = False,
        summarize_images: bool = False,
    ):
        self.cfg = cfg
        self.model_config = model_config
        self.enable_gmft = enable_gmft
        self.do_ocr = do_ocr
        self.parse_tables = parse_tables
        self.summarize_images = summarize_images
        # Support both dict and OmegaConf access patterns
        text_config = model_config.get('text') if isinstance(model_config, dict) else getattr(model_config, 'text', None)
        vision_config = model_config.get('vision') if isinstance(model_config, dict) else getattr(model_config, 'vision', None)
        self.table_summarizer = TableSummarizer(self.cfg, text_config) if self.parse_tables and text_config else None
        self.image_summarizer = ImageSummarizer(self.cfg, vision_config) if self.summarize_images and vision_config else None
        self.verbose = verbose

        # Parallel summarization settings
        try:
            self.summarization_workers = cfg.doc_processing.get("summarization_workers", 4)
        except (AttributeError, KeyError):
            self.summarization_workers = 4
        try:
            self.store_image_bytes = cfg.doc_processing.get("add_image_bytes", False)
        except (AttributeError, KeyError):
            self.store_image_bytes = False

    def cleanup(self):
        """Release heavy resources. Subclasses may override for parser-specific cleanup."""
        self.table_summarizer = None
        self.image_summarizer = None

    def _parallel_summarize_images(self, tasks):
        """
        Summarize images, optionally in parallel via ThreadPoolExecutor.

        Args:
            tasks: List of dicts with keys:
                image_path, source_url, previous_text, next_text, image_bytes

        Returns:
            List of summary strings (or None for failures), preserving input order.
        """
        if not tasks:
            return []

        if not self.image_summarizer:
            logger.warning("_parallel_summarize_images called but image_summarizer is None")
            return [None] * len(tasks)

        def _summarize_one(task):
            try:
                return self.image_summarizer.summarize_image(
                    task['image_path'],
                    task['source_url'],
                    task.get('previous_text'),
                    task.get('next_text'),
                    image_bytes=task.get('image_bytes')
                )
            except Exception as e:
                logger.error(f"Image summarization failed: {e}")
                return None

        if self.summarization_workers <= 1:
            return [_summarize_one(t) for t in tasks]

        results = [None] * len(tasks)
        with ThreadPoolExecutor(max_workers=self.summarization_workers) as executor:
            future_to_idx = {
                executor.submit(_summarize_one, task): idx
                for idx, task in enumerate(tasks)
            }
            for future in as_completed(future_to_idx):
                idx = future_to_idx[future]
                try:
                    results[idx] = future.result()
                except Exception as e:
                    logger.error(f"Image summarization future failed for task {idx}: {e}")
        return results

    def _parallel_summarize_tables(self, table_texts):
        """
        Summarize tables, optionally in parallel via ThreadPoolExecutor.

        Args:
            table_texts: List of table text/markdown strings.

        Returns:
            List of summary strings (or "" for failures), preserving input order.
        """
        if not table_texts:
            return []

        if not self.table_summarizer:
            logger.warning("_parallel_summarize_tables called but table_summarizer is None")
            return [""] * len(table_texts)

        def _summarize_one(text):
            try:
                return self.table_summarizer.summarize_table_text(text)
            except Exception as e:
                logger.error(f"Table summarization failed: {e}")
                return ""

        if self.summarization_workers <= 1:
            return [_summarize_one(t) for t in table_texts]

        results = [""] * len(table_texts)
        with ThreadPoolExecutor(max_workers=self.summarization_workers) as executor:
            future_to_idx = {
                executor.submit(_summarize_one, text): idx
                for idx, text in enumerate(table_texts)
            }
            for future in as_completed(future_to_idx):
                idx = future_to_idx[future]
                try:
                    results[idx] = future.result()
                except Exception as e:
                    logger.error(f"Table summarization future failed for task {idx}: {e}")
        return results

    def parse(self, filename: str, source_url: str = "No URL") -> ParsedDocument:
        """
        Parse a document and return unified content stream with proper element ordering.
        Subclasses must override this method.

        Args:
            filename: Path to the file to parse
            source_url: Source URL for context

        Returns:
            ParsedDocument with unified content stream
        """
        raise NotImplementedError("Subclasses must implement parse() method")

    def get_tables_with_gmft(self, filename: str) -> List[Tuple[pd.DataFrame, str, str, dict]]:
        if not filename.endswith('.pdf'):
            logger.warning(f"GMFT: only PDF files are supported, skipping {filename}")
            return []

        detector = TableDetector()
        config = AutoFormatConfig()
        config.semantic_spanning_cells = True   # [Experimental] better spanning cells
        config.enable_multi_header = True       # multi-indices
        formatter = AutoTableFormatter(config)

        table_data = []  # (df, markdown, captions, metadata)
        doc = PyPDFium2Document(filename)
        try:
            for inx, page in enumerate(doc):
                if inx % 100 == 0:
                    logger.info(f"GMFT: processing page {inx+1}")
                for table in detector.extract(page):
                    try:
                        ft = formatter.extract(table)
                        df = ft.df()
                        table_data.append((df, df.to_markdown(), " ".join(ft.captions()), {'page': inx+1}))
                    except Exception as e:
                        logger.error(f"Error processing table (with GMFT) on page {inx+1}: {e}")
                        continue
        finally:
            doc.close()
            if self.verbose:
                logger.info("GMFT: Finished processing PDF")

        if not table_data:
            return []

        table_texts = [md for _, md, _, _ in table_data]
        summaries = self._parallel_summarize_tables(table_texts)
        return [(df, summary, captions, metadata)
                for (df, _, captions, metadata), summary in zip(table_data, summaries)]


class DocupandaDocumentParser(DocumentParser):

    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = {},
        docupanda_api_key: str = None,
        parse_tables: bool = False,
        enable_gmft: bool = False,
        summarize_images: bool = False,
        image_context: dict = None
    ):
        super().__init__(
            verbose=verbose,
            cfg=cfg,
            model_config=model_config,
            parse_tables=parse_tables,
            enable_gmft=enable_gmft,
            do_ocr=False,
            summarize_images=summarize_images
        )
        self._api_key = docupanda_api_key
        self.image_context = image_context or {'num_previous_chunks': 1, 'num_next_chunks': 1}
        if not self._api_key:
            raise ValueError("No Docupanda API key found, skipping Docupanda")

    def parse(
            self,
            filename: str,
            source_url: str = "No URL"
        ) -> ParsedDocument:
        """
        Parse a document and return unified content stream with images interleaved in proper order.
        Tables are extracted separately for structured indexing.
        Using DocuPanda

        Args:
            filename (str): Name of the file to parse.
            source_url (str): Source URL for context.

        Returns:
            ParsedDocument with unified content stream
        """
        st = time.time()
        doc_title = extract_document_title(filename)
        
        all_elements = []  # For building unified content stream
        tables = []
        image_bytes = []  # Store image binary data
        img_folder = '/images'
        base_url = 'https://app.docupanda.io/document'
        os.makedirs(img_folder, exist_ok=True)

        # Phase 1: post document
        payload = {"document": {"file": {
            "contents": base64.b64encode(open(filename, 'rb').read()).decode(),
                "filename": filename
            }}}
        headers = {
            "accept": "application/json",
            "content-type": "application/json",
            "X-API-Key": self._api_key
        }
        response = requests.post(base_url, json=payload, headers=headers)
        if response.status_code != 200:
            logger.error(f"Docupanda: failed to post document, response code {response.status_code}, msg={response.json()}")
            return ParsedDocument(title='', content_stream=[], tables=[], image_bytes=[])

        # Phase 2: get document; poll until ready, but no longer than timeout
        document_id = response.json()['documentId']
        timeout = 60
        start_time = time.time()
        completed = False
        while time.time() - start_time < timeout:
            response = requests.get(f"{base_url}/{document_id}", headers=headers)
            if response.json()['status'] == 'completed':
                completed = True
                break
            time.sleep(1)

        if not completed:
            logger.error(f"Docupanda: document processing timed out after {timeout} seconds")
            return ParsedDocument(title='', content_stream=[], tables=[], image_bytes=[])

        # Phase 3: get results
        pdf_bytes = pathlib.Path(filename).read_bytes()
        extracted_images = convert_from_bytes(pdf_bytes, dpi=300)
        img_num = 0

        # First pass: collect all sections for context extraction
        all_sections = []
        element_index = 0
        for page in response.json()['result']['pages']:
            page_num = page['pageNum']
            base_position = page_num * 1000
            
            for section in page['sections']:
                position = base_position + element_index
                all_sections.append({
                    'section': section,
                    'page_num': page_num,
                    'position': position,
                    'element_index': element_index
                })
                element_index += 1

        # Second pass: collect text, tables, and image tasks
        image_tasks = []
        table_data = []  # (df, markdown, page_num)
        for idx, section_info in enumerate(all_sections):
            section = section_info['section']
            page_num = section_info['page_num']
            position = section_info['position']

            if section['type'] == 'text':
                metadata = {
                    'element_type': 'text',
                    'page': page_num
                }
                all_elements.append((position, section['text'], metadata))

            elif section['type'] == 'table':
                tableList = section['tableList']
                header = tableList[0]
                rows = tableList[1:]
                df = pd.DataFrame(rows, columns=header)
                table_data.append((df, df.to_markdown(), page_num))
                if self.verbose:
                    logger.info(f"Table found on page {page_num} - will be processed for structured indexing")

            elif section['type'] == 'image' and self.summarize_images:
                previous_text, next_text = extract_image_context(
                    [s['section'] for s in all_sections],
                    idx,
                    num_previous=self.image_context['num_previous_chunks'],
                    num_next=self.image_context['num_next_chunks'],
                    text_extractor=lambda x: x.get('text') if x.get('type') == 'text' else None
                )

                bbox = [round(x, 2) for x in section['bbox']]
                page_idx = page_num - 1
                if page_idx < 0 or page_idx >= len(extracted_images):
                    logger.warning(f"Docupanda: page_num {page_num} out of range for rendered pages ({len(extracted_images)}), skipping image")
                    img_num += 1
                    continue
                img = extracted_images[page_idx]
                image_dims = img.size
                bbox_pixels = (int(bbox[0] * image_dims[0]), int(bbox[1] * image_dims[1]),
                               int(bbox[2] * image_dims[0]), int(bbox[3] * image_dims[1]))
                img_cropped = img.crop(box=bbox_pixels)

                buf = BytesIO()
                img_cropped.save(buf, format='PNG')
                image_binary = buf.getvalue()
                buf.close()
                del img_cropped

                image_id = f"docupanda_page_{page_num}_image_{img_num}"
                image_tasks.append({
                    'image_path': '', 'source_url': source_url,
                    'previous_text': previous_text, 'next_text': next_text,
                    'image_bytes': image_binary,
                    'image_id': image_id, 'position': position, 'page_no': page_num,
                    'bbox': bbox,
                })
                img_num += 1

            else:
                logger.info(f"Docupanda: unknown section type {section['type']} on page {page_num} ignored...")

        # Free rendered pages now that image crops are extracted
        del extracted_images

        # Parallel table summarization
        if table_data:
            table_texts = [md for _, md, _ in table_data]
            table_summaries = self._parallel_summarize_tables(table_texts)
            for (df, _, page_num), summary in zip(table_data, table_summaries):
                summary = summary if summary else df.to_markdown()
                tables.append([df, summary, '', {'page': page_num}])

        # Parallel image summarization
        if image_tasks:
            image_summaries = self._parallel_summarize_images(image_tasks)
            for task, summary in zip(image_tasks, image_summaries):
                if self.store_image_bytes:
                    image_bytes.append((task['image_id'], task['image_bytes']))
                if summary and len(summary) > 10:
                    metadata = {
                        'element_type': 'image',
                        'page': task['page_no'],
                        'image_id': task['image_id']
                    }
                    all_elements.append((task['position'] + 0.5, summary, metadata))
                    if self.verbose:
                        logger.info(f"Image summary: {summary[:MAX_VERBOSE_LENGTH]}...")
                logger.info(f"Docupanda: processed image with bounding box {task['bbox']} on page {task['page_no']}")
                # Release binary memory
                task['image_bytes'] = None

        # Sort all elements by position to maintain document order
        all_elements.sort(key=lambda x: x[0])
        content_stream = [(content, metadata) for _, content, metadata in all_elements]

        # Fallback to filename if no title found
        if not doc_title:
            basename = os.path.basename(filename)
            doc_title = os.path.splitext(basename)[0].replace('_', ' ').replace('-', ' ').title()
            logger.info(f"No title found in document, using filename fallback: '{doc_title}' for file {filename}")

        logger.info(f"DocupandaParser unified: {len(content_stream)} content elements, {len(tables)} tables")
        logger.info(f"parsing file {filename} with Docupanda took {time.time()-st:.2f} seconds")

        return ParsedDocument(
            title=doc_title,
            content_stream=content_stream,
            tables=tables,
            image_bytes=image_bytes
        )

class LlamaParseDocumentParser(DocumentParser):

    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = {},
        llama_parse_api_key: str = None,
        parse_tables: bool = False,
        enable_gmft: bool = False,
        summarize_images: bool = False,
        image_context: dict = None
    ):
        super().__init__(
            cfg=cfg,
            verbose=verbose,
            model_config=model_config,
            parse_tables=parse_tables,
            enable_gmft=enable_gmft,
            do_ocr=False,
            summarize_images=summarize_images
        )
        self.image_context = image_context or {'num_previous_chunks': 1, 'num_next_chunks': 1}
        if llama_parse_api_key:
            self.parser = LlamaParse(verbose=True, premium_mode=True, api_key=llama_parse_api_key)
            if self.verbose:
                logger.info("Using LlamaParse, premium mode")
        else:
            raise ValueError("No LlamaParse API key found, skipping LlamaParse")

    def parse(self, filename: str, source_url: str = "No URL") -> ParsedDocument:
        """
        Parse a document and return unified content stream with images interleaved in proper order.
        Tables are extracted separately for structured indexing.
        """
        st = time.time()
        doc_title = extract_document_title(filename)
        
        image_bytes = []  # Store image binary data
        img_folder = '/images'
        os.makedirs(img_folder, exist_ok=True)

        nest_asyncio.apply()
        json_objs = self.parser.get_json_result(filename)
        
        # Collect all elements with their positions for proper ordering
        all_elements = []
        
        # First pass: collect all text elements per page
        page_elements = {}  # page_num -> list of (element_type, content, metadata)
        
        for page_data in json_objs[0]['pages']:
            page_num = page_data['page']
            if page_num not in page_elements:
                page_elements[page_num] = []
            
            # Add page text
            page_text = page_data['text']
            if page_text and page_text.strip():
                metadata = {
                    'element_type': 'text', 
                    'page': page_num
                }
                page_elements[page_num].append(('text', page_text, metadata))
            
            # Tables are not added inline - they are processed separately for structured indexing
        
        # Second pass: collect image tasks for parallel summarization
        if self.summarize_images:
            parsed_images = self.parser.get_images(json_objs, download_path=img_folder)

            flat_elements = []
            for page_num in sorted(page_elements.keys()):
                for elem in page_elements[page_num]:
                    flat_elements.append(elem)

            image_tasks = []
            image_counter = 0
            for image_dict in parsed_images:
                page_num = image_dict['page_number']

                image_idx = len(flat_elements)
                for idx, (elem_type, content, meta) in enumerate(flat_elements):
                    if meta['page'] == page_num:
                        image_idx = idx + 1
                        break

                previous_text, next_text = extract_image_context(
                    flat_elements, image_idx,
                    num_previous=self.image_context['num_previous_chunks'],
                    num_next=self.image_context['num_next_chunks'],
                    text_extractor=lambda x: x[1] if x[0] == 'text' else None
                )

                image_binary = None
                if self.store_image_bytes:
                    with open(image_dict['path'], 'rb') as fp:
                        image_binary = fp.read()
                image_id = f"llamaparse_page_{page_num}_image_{image_counter}"
                image_counter += 1

                image_tasks.append({
                    'image_path': image_dict['path'], 'source_url': source_url,
                    'previous_text': previous_text, 'next_text': next_text,
                    'image_bytes': image_binary,
                    'image_id': image_id, 'page_num': page_num,
                })

            # Parallel image summarization
            if image_tasks:
                summaries = self._parallel_summarize_images(image_tasks)
                for task, summary in zip(image_tasks, summaries):
                    if self.store_image_bytes and task['image_bytes']:
                        image_bytes.append((task['image_id'], task['image_bytes']))
                    if summary:
                        page_num = task['page_num']
                        if page_num not in page_elements:
                            page_elements[page_num] = []
                        metadata = {
                            'element_type': 'image',
                            'page': page_num,
                            'image_id': task['image_id']
                        }
                        page_elements[page_num].append(('image', summary, metadata))
                    task['image_bytes'] = None  # Release memory
        
        # Now build the final positioned elements using standard formula
        for page_num in sorted(page_elements.keys()):
            base_position = page_num * 1000
            for element_index, (element_type, content, metadata) in enumerate(page_elements[page_num]):
                if element_type == 'image':
                    position = base_position + element_index + 0.5
                else:
                    position = base_position + element_index
                all_elements.append((position, content, metadata))
                
                if self.verbose:
                    if element_type == 'image':
                        logger.info(f"Image summary: {content[:MAX_VERBOSE_LENGTH]}...")

        # Sort all elements by position to maintain document order
        all_elements.sort(key=lambda x: x[0])
        content_stream = [(content, metadata) for _, content, metadata in all_elements]

        # Process tables separately for structured indexing
        tables = []
        if self.parse_tables:
            if self.enable_gmft and filename.endswith('.pdf'):
                tables = list(self.get_tables_with_gmft(filename))
            else:
                # Collect table data, then parallel-summarize
                lp_table_data = []  # (df, markdown, page_num)
                for page_data in json_objs[0]['pages']:
                    page_num = page_data['page']
                    if 'items' in page_data:
                        for item in page_data['items']:
                            if item['type'] == 'table':
                                table_md = item['md']
                                lp_table_data.append((markdown_to_df(table_md), table_md, page_num))

                if lp_table_data and self.table_summarizer:
                    table_texts = [md for _, md, _ in lp_table_data]
                    table_summaries = self._parallel_summarize_tables(table_texts)
                    for (df, md, page_num), summary in zip(lp_table_data, table_summaries):
                        tables.append([df, summary if summary else md, '', {'page': page_num}])
                elif lp_table_data:
                    for df, md, page_num in lp_table_data:
                        tables.append([df, md, '', {'page': page_num}])

        # Fallback to filename if no title found
        if not doc_title:
            basename = os.path.basename(filename)
            doc_title = os.path.splitext(basename)[0].replace('_', ' ').replace('-', ' ').title()
            logger.info(f"No title found in document, using filename fallback: '{doc_title}' for file {filename}")

        logger.info(f"LlamaParseParser: {len(content_stream)} content elements, {len(tables)} tables")
        logger.info(f"parsing file {filename} with LlamaParse took {time.time()-st:.2f} seconds")

        return ParsedDocument(
            title=doc_title,
            content_stream=content_stream,
            tables=tables,
            image_bytes=image_bytes
        )

class DoclingDocumentParser(DocumentParser):

    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = {},
        chunking_strategy: str = 'hierarchical',
        chunk_size: int = 1024,
        parse_tables: bool = False,
        enable_gmft: bool = False,
        do_ocr: bool = False,
        summarize_images: bool = False,
        image_scale: float = 1.0,
        image_context: dict = None,
        layout_model: str = None,
        do_formula_enrichment: bool = False,
        fallback_ocr: bool = False,
        pdf_batch_size: int = 500
    ):
        super().__init__(
            cfg=cfg,
            verbose=verbose,
            model_config=model_config,
            parse_tables=parse_tables,
            enable_gmft=enable_gmft,
            do_ocr=do_ocr,
            summarize_images=summarize_images
        )
        self.chunking_strategy = chunking_strategy
        self.chunk_size = chunk_size
        self.pdf_batch_size = pdf_batch_size
        self.image_scale = image_scale
        self.image_context = image_context or {'num_previous_chunks': 1, 'num_next_chunks': 1}
        self.layout_model = layout_model  # Options: None (default heron), 'heron', 'heron_101', 'v2'
        self.do_formula_enrichment = do_formula_enrichment
        self.fallback_ocr = fallback_ocr
        self._converter = None
        self._ocr_converter = None
        if self.verbose:
            layout_info = f", layout_model '{self.layout_model}'" if self.layout_model else ""
            logger.info(f"Using DoclingParser with chunking strategy {self.chunking_strategy} and chunk size {self.chunk_size}{layout_info}")

    @staticmethod
    def _lazy_load_docling():
        import warnings

        warnings.filterwarnings(
            "ignore",
            message="Token indices sequence length",
            category=UserWarning
        )

        from docling.document_converter import DocumentConverter, PdfFormatOption, HTMLFormatOption, PowerpointFormatOption, WordFormatOption
        from docling.datamodel.pipeline_options import PdfPipelineOptions, EasyOcrOptions, RapidOcrOptions, PaginatedPipelineOptions
        from docling.datamodel.base_models import InputFormat
        from docling_core.transforms.chunker.hybrid_chunker import HybridChunker
        from docling_core.transforms.chunker import HierarchicalChunker
        from docling.datamodel.pipeline_options import LayoutOptions
        from docling.datamodel.layout_model_specs import DOCLING_LAYOUT_HERON, DOCLING_LAYOUT_HERON_101, DOCLING_LAYOUT_V2

        layout_model_specs = {}
        layout_model_specs['LayoutOptions'] = LayoutOptions
        layout_model_specs['heron'] = DOCLING_LAYOUT_HERON
        layout_model_specs['heron_101'] = DOCLING_LAYOUT_HERON_101
        layout_model_specs['v2'] = DOCLING_LAYOUT_V2

        return (
            DocumentConverter, HybridChunker, HierarchicalChunker, PdfPipelineOptions,
            PdfFormatOption, HTMLFormatOption, PowerpointFormatOption, WordFormatOption, InputFormat,
            EasyOcrOptions, RapidOcrOptions, PaginatedPipelineOptions, layout_model_specs
        )

    def _get_tables(self, tables, doc=None):
        def _get_metadata(table):
            md = {}
            if table.prov:
                md['page'] = table.prov[0].page_no
            return md

        table_data = []
        for table in tables:
            try:
                table_df = table.export_to_dataframe(doc) if doc else table.export_to_dataframe()
                table_md = table_df.to_markdown()
                table_data.append((table_df, table_md, _get_metadata(table)))
            except ValueError as err:
                logger.error(f"Error parsing Markdown table: {err}. Skipping...")
                continue

        if not table_data:
            return []

        table_texts = [md for _, md, _ in table_data]
        summaries = self._parallel_summarize_tables(table_texts)
        return [(df, summary, '', metadata)
                for (df, _, metadata), summary in zip(table_data, summaries)]

    def cleanup(self):
        """Release the cached Docling converter and its ML models."""
        super().cleanup()
        self._converter = None
        self._ocr_converter = None

    @staticmethod
    def _get_pdf_page_count(filename: str) -> int:
        from pypdf import PdfReader
        return len(PdfReader(filename).pages)

    def _get_or_create_converter(self, fallback_ocr: bool = False):
        """Build the DocumentConverter once and cache it for reuse across files.

        Args:
            fallback_ocr: When True, returns a separate OCR-enabled converter
                       cached in self._ocr_converter.
        """
        if fallback_ocr:
            if self._ocr_converter is not None:
                return self._ocr_converter
        else:
            if self._converter is not None:
                return self._converter

        (
            DocumentConverter, _, _, PdfPipelineOptions,
            PdfFormatOption, HTMLFormatOption, PowerpointFormatOption, WordFormatOption, InputFormat,
            EasyOcrOptions, RapidOcrOptions, PaginatedPipelineOptions, layout_model_specs
        ) = self._lazy_load_docling()

        from docling.backend.html_backend import HTMLDocumentBackend

        request_headers = get_headers(self.cfg)
        source_url_holder: dict = {'url': ''}

        class UrlCapturingHtmlBackend(HTMLDocumentBackend):
            """Fetches remote images with browser headers and correctly resolves
            root-relative URLs (/path) that Docling's default resolver mishandles.

            Docling's _resolve_relative_path treats /path as a protocol-relative URL
            and produces "https:/path" (one slash, no domain). This override fixes
            that by reconstructing the origin from base_path.
            """

            def __init__(self, in_doc, path_or_stream, options=None):
                if options is None:
                    from docling.datamodel.backend_options import HTMLBackendOptions as _HBOpts
                    options = _HBOpts()
                super().__init__(in_doc, path_or_stream, options)
                url = source_url_holder.get('url', '')
                if url and url.startswith(('http://', 'https://')):
                    self.base_path = url

            def _resolve_relative_path(self, loc: str) -> str:
                # Docling treats root-relative /path as protocol-relative //host/path,
                # producing "https:/path" (single slash). Fix: prepend scheme+netloc.
                if self.base_path and loc.startswith('/') and not loc.startswith('//'):
                    parsed = urlparse(self.base_path)
                    if parsed.scheme and parsed.netloc:
                        return f"{parsed.scheme}://{parsed.netloc}{loc}"
                return super()._resolve_relative_path(loc)

            def _load_image_data(self, src_loc: str) -> Optional[bytes]:
                logger.debug(f"_load_image_data: src_loc={src_loc!r}")
                if HTMLDocumentBackend._is_remote_url(src_loc):
                    try:
                        r = requests.get(src_loc, stream=True, headers=request_headers, timeout=10)
                        r.raise_for_status()
                        if r.headers.get("content-type", "").startswith("image/"):
                            from PIL import Image as _PILImage
                            data = r.content
                            # Validate completeness — PIL loads lazily; a truncated image
                            # only fails during save(), which Docling doesn't catch.
                            _PILImage.open(BytesIO(data)).load()
                            return data
                    except Exception as e:
                        logger.debug(f"Failed to fetch image from {src_loc}: {e}")
                    return None
                return super()._load_image_data(src_loc)

        html_opts = PaginatedPipelineOptions()
        html_opts.generate_page_images = False  # Avoid page-render crops; use URL fetch instead
        html_opts.images_scale = self.image_scale

        pdf_opts = PdfPipelineOptions()
        pdf_opts.images_scale = self.image_scale
        pdf_opts.generate_picture_images = True
        pdf_opts.do_ocr = False
        pdf_opts.do_formula_enrichment = self.do_formula_enrichment

        # Configure layout model if specified
        valid_layout_models = [k for k in layout_model_specs.keys() if k != 'LayoutOptions']
        if self.layout_model and self.layout_model in valid_layout_models:
            LayoutOptions = layout_model_specs.get('LayoutOptions')
            model_spec = layout_model_specs.get(self.layout_model)
            if LayoutOptions and model_spec:
                pdf_opts.layout_options = LayoutOptions(model_spec=model_spec)
                logger.info(f"Using custom layout model: {self.layout_model}")
        elif self.layout_model:
            logger.warning(f"Layout model '{self.layout_model}' not found. Available: {[k for k in layout_model_specs.keys() if k != 'LayoutOptions']}")

        # Pipeline options for Office documents
        office_opts = PaginatedPipelineOptions()
        office_opts.generate_page_images = True
        office_opts.images_scale = self.image_scale

        # Get OCR engine preference from config (default to 'easyocr')
        ocr_engine = getattr(self.cfg.doc_processing, "ocr_engine", "easyocr")

        # Configure OCR options based on engine choice
        if ocr_engine == "rapidocr":
            ocr_config = getattr(self.cfg.doc_processing, "rapid_ocr_config", None)
            ocr_config = OmegaConf.to_container(
                ocr_config, resolve=True
            ) if ocr_config is not None else {}
            ocr_options = RapidOcrOptions()
            rapid_mapping = {
                "bitmap_area_threshold": lambda val: setattr(ocr_options, "bitmap_area_threshold", val),
                "force_full_page_ocr": lambda val: setattr(ocr_options, "force_full_page_ocr", val),
                "lang": lambda val: setattr(ocr_options, "lang", val),
                "use_det": lambda val: setattr(ocr_options, "use_det", val),
                "use_cls": lambda val: setattr(ocr_options, "use_cls", val),
                "text_score": lambda val: setattr(ocr_options, "text_score", val),
            }
            for key, func in rapid_mapping.items():
                value = ocr_config.get(key)
                if value is not None:
                    func(value)
            logger.info("Using RapidOCR engine for PDF processing")
        else:  # Default to EasyOCR
            ocr_config = getattr(self.cfg.doc_processing, "easy_ocr_config", None)
            ocr_config = OmegaConf.to_container(
                ocr_config, resolve=True
            ) if ocr_config is not None else {}
            ocr_options = EasyOcrOptions()
            easy_mapping = {
                "bitmap_area_threshold": lambda val: setattr(ocr_options, "bitmap_area_threshold", val),
                "confidence_threshold": lambda val: setattr(ocr_options, "confidence_threshold", val),
                "download_enabled": lambda val: setattr(ocr_options, "download_enabled", val),
                "force_full_page_ocr": lambda val: setattr(ocr_options, "force_full_page_ocr", val),
                "kind": lambda val: setattr(ocr_options, "kind", val),
                "lang": lambda val: setattr(ocr_options, "lang", val),
                "model_config": lambda val: setattr(ocr_options, "model_config", val),
                "model_storage_directory": lambda val: setattr(ocr_options, "model_storage_directory", val),
                "recog_network": lambda val: setattr(ocr_options, "recog_network", val),
                "use_gpu": lambda val: setattr(ocr_options, "use_gpu", val),
            }
            for key, func in easy_mapping.items():
                value = ocr_config.get(key)
                if value is not None:
                    func(value)
            logger.info("Using EasyOCR engine for PDF processing")

        # Always set OCR options to prevent pipeline initialization failure
        pdf_opts.ocr_options = ocr_options

        if self.do_ocr or fallback_ocr:
            pdf_opts.do_ocr = True

        from docling.datamodel.backend_options import HTMLBackendOptions
        html_backend_opts = HTMLBackendOptions(fetch_images=True, enable_remote_fetch=True)

        converter = DocumentConverter(
            allowed_formats=[
                InputFormat.PDF, InputFormat.HTML, InputFormat.PPTX, InputFormat.DOCX,
            ],
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=pdf_opts),
                InputFormat.HTML: HTMLFormatOption(
                    backend=UrlCapturingHtmlBackend,
                    pipeline_options=html_opts,
                    backend_options=html_backend_opts,
                ),
                InputFormat.PPTX: PowerpointFormatOption(pipeline_options=office_opts),
                InputFormat.DOCX: WordFormatOption(pipeline_options=office_opts),
            }
        )
        if fallback_ocr:
            self._ocr_converter = converter
        else:
            self._converter = converter
            self._source_url_holder = source_url_holder
        return converter

    def _extract_from_doc(self, doc, source_url, image_counter_start=0):
        """
        Extract text elements, image tasks, and tables from a single Docling document.

        Returns (positioned_elements, image_tasks, tables, next_image_counter).
        """
        positioned_elements = []
        image_tasks = []
        image_counter = image_counter_start

        # First pass: collect all items for context extraction
        all_items = []
        element_index = 0
        for item, _ in doc.iterate_items():
            page_no = 0
            if hasattr(item, 'prov') and item.prov:
                page_no = item.prov[0].page_no

            base_position = page_no * 1000
            position = base_position + element_index

            all_items.append({
                'item': item,
                'position': position,
                'page_no': page_no,
                'index': element_index
            })
            element_index += 1

        items_list = [i['item'] for i in all_items]

        # Second pass: process text inline, collect image tasks for parallel summarization
        for idx, item_info in enumerate(all_items):
            item = item_info['item']
            page_no = item_info['page_no']
            position = item_info['position']

            if hasattr(item, 'export_to_dataframe'):
                if self.verbose and self.parse_tables:
                    logger.info(f"Table found on page {page_no} - will be processed for structured indexing")

            elif hasattr(item, 'text'):
                metadata = {
                    'element_type': 'text',
                    'page': page_no
                }
                positioned_elements.append((position, item.text, metadata))

            elif hasattr(item, 'get_image') and self.summarize_images:
                previous_text, next_text = extract_image_context(
                    items_list, idx,
                    num_previous=self.image_context['num_previous_chunks'],
                    num_next=self.image_context['num_next_chunks'],
                    text_extractor=lambda x: x.text if hasattr(x, 'text') else None
                )

                image_binary = None
                image = item.get_image(doc)
                if image:
                    w, h = image.size
                    if min(w, h) < MIN_IMAGE_DIMENSION:
                        logger.debug(f"Skipping small image ({w}×{h}px) — likely icon/avatar")
                    else:
                        buf = BytesIO()
                        image.save(buf, format='PNG')
                        image_binary = buf.getvalue()
                        buf.close()
                        logger.debug(f"Image {image_counter}: {w}×{h}px from docling fetch")
                    del image

                if image_binary:
                    image_id = f"docling_page_{page_no}_image_{image_counter}"
                    image_counter += 1
                    image_tasks.append({
                        'image_path': '', 'source_url': source_url,
                        'previous_text': previous_text, 'next_text': next_text,
                        'image_bytes': image_binary,
                        'image_id': image_id, 'position': position, 'page_no': page_no,
                    })
                else:
                    logger.debug("Could not retrieve image (no embedded data and no fetchable URL)")

        # Process tables (needs doc.tables)
        tables = []
        if self.parse_tables and not (self.enable_gmft and self._current_filename.endswith('.pdf')):
            tables = list(self._get_tables(doc.tables, doc))

        return positioned_elements, image_tasks, tables, image_counter

    def _apply_chunking(self, doc, positioned_elements, HybridChunker, HierarchicalChunker):
        """Apply chunking to text elements using the Docling doc object."""
        chunker = (
            HybridChunker(max_tokens=self.chunk_size)
            if self.chunking_strategy == 'hybrid' else HierarchicalChunker()
        )

        non_text_elements = [(pos, content, meta) for pos, content, meta in positioned_elements
                           if meta['element_type'] != 'text']

        chunked_text_elements = []
        for chunk in chunker.chunk(doc):
            metadata = {'element_type': 'text'}
            if chunk.meta.doc_items and chunk.meta.doc_items[0].prov:
                page_no = chunk.meta.doc_items[0].prov[0].page_no
                metadata['page'] = page_no
            else:
                metadata['page'] = 0

            base_position = metadata['page'] * 1000
            position = base_position + len([e for e in chunked_text_elements if e[2]['page'] == metadata['page']])
            chunked_text_elements.append((position, chunker.serialize(chunk=chunk), metadata))

        return chunked_text_elements + non_text_elements

    def _needs_batching(self, filename: str) -> bool:
        """Check if a PDF file needs batch processing."""
        if not filename.lower().endswith('.pdf'):
            return False
        try:
            page_count = self._get_pdf_page_count(filename)
            return page_count > self.pdf_batch_size
        except Exception as e:
            logger.warning(f"Could not determine page count for {filename}: {e}")
            return False

    def _parse_with_converter(self, filename: str, source_url: str, converter) -> ParsedDocument:
        """
        Core parsing logic shared by the normal and OCR-fallback paths.

        Converts *filename* with *converter*, extracts text / images / tables,
        applies chunking, and returns a ``ParsedDocument``.

        For large PDFs (page count > pdf_batch_size), processes in batches
        to avoid OOM kills from Docling's ML pipeline.
        """
        (
            _, HybridChunker, HierarchicalChunker, _,
            _, _, _, _, _,
            _, _, _, _
        ) = self._lazy_load_docling()

        if hasattr(self, '_source_url_holder'):
            self._source_url_holder['url'] = source_url

        # Store filename for _extract_from_doc to check gmft eligibility
        self._current_filename = filename

        doc_title = extract_document_title(filename)

        if self._needs_batching(filename):
            return self._parse_batched(
                filename, source_url, converter, doc_title,
                HybridChunker, HierarchicalChunker
            )

        # --- Single-pass (non-batched) path ---
        res = converter.convert(filename)
        doc = res.document

        if not doc_title and doc.name:
            doc_title = doc.name
            logger.info(f"Using Docling document name: '{doc_title}' from file {filename}")

        positioned_elements, image_tasks, tables, _ = self._extract_from_doc(doc, source_url)

        # Apply chunking BEFORE image summarization (chunker needs doc object)
        if self.chunking_strategy in ['hybrid', 'hierarchical']:
            positioned_elements = self._apply_chunking(
                doc, positioned_elements, HybridChunker, HierarchicalChunker
            )

        # Free the heavy Docling document ASAP — before the slow image
        # summarization API calls, which can keep it alive for minutes.
        del doc, res
        gc.collect()
        release_memory()

        # GMFT tables (needs the file on disk, not the doc object)
        if self.parse_tables and self.enable_gmft and filename.endswith('.pdf'):
            tables = list(self.get_tables_with_gmft(filename))

        # Parallel image summarization (runs after doc is freed)
        image_bytes = []
        if image_tasks:
            image_summaries = self._parallel_summarize_images(image_tasks)
            for task, summary in zip(image_tasks, image_summaries):
                if self.store_image_bytes:
                    image_bytes.append((task['image_id'], task['image_bytes']))
                if summary:
                    metadata = {
                        'element_type': 'image',
                        'page': task['page_no'],
                        'image_id': task['image_id']
                    }
                    positioned_elements.append((task['position'] + 0.5, summary, metadata))
                    if self.verbose:
                        logger.info(f"Image summary at position {task['position'] + 0.5}: {summary[:MAX_VERBOSE_LENGTH]}...")
                task['image_bytes'] = None

        return self._finalize(filename, doc_title, positioned_elements, tables, image_bytes)

    def _parse_batched(self, filename, source_url, converter, doc_title,
                       HybridChunker, HierarchicalChunker):
        """Process a large PDF in page-range batches to limit memory usage."""
        total_pages = self._get_pdf_page_count(filename)
        batch_size = self.pdf_batch_size
        logger.info(f"Batching {filename}: {total_pages} pages in batches of {batch_size}")

        all_positioned = []
        all_tables = []
        image_counter = 0

        # Spill image bytes to a temp directory so they don't accumulate in RAM
        # alongside the docling ML models during batch processing.
        img_tmp_dir = tempfile.mkdtemp(prefix="vectara_img_") if self.store_image_bytes else None
        img_manifest = []  # list of (image_id, file_path)

        for batch_start in range(1, total_pages + 1, batch_size):
            batch_end = min(batch_start + batch_size - 1, total_pages)
            logger.info(f"Processing batch pages {batch_start}-{batch_end} of {total_pages}")

            res = converter.convert(filename, page_range=(batch_start, batch_end))
            doc = res.document

            if not doc_title and doc.name:
                doc_title = doc.name

            positioned_elements, image_tasks, tables, image_counter = self._extract_from_doc(
                doc, source_url, image_counter_start=image_counter
            )

            # Apply chunking per-batch (chunker needs doc object)
            if self.chunking_strategy in ['hybrid', 'hierarchical']:
                positioned_elements = self._apply_chunking(
                    doc, positioned_elements, HybridChunker, HierarchicalChunker
                )

            all_positioned.extend(positioned_elements)
            all_tables.extend(tables)

            # Summarize images per-batch to avoid accumulating image bytes across all batches
            if image_tasks:
                image_summaries = self._parallel_summarize_images(image_tasks)
                for task, summary in zip(image_tasks, image_summaries):
                    if self.store_image_bytes and task['image_bytes']:
                        img_path = os.path.join(img_tmp_dir, f"{task['image_id']}.bin")
                        with open(img_path, 'wb') as f:
                            f.write(task['image_bytes'])
                        img_manifest.append((task['image_id'], img_path))
                    if summary:
                        metadata = {
                            'element_type': 'image',
                            'page': task['page_no'],
                            'image_id': task['image_id']
                        }
                        all_positioned.append((task['position'] + 0.5, summary, metadata))
                        if self.verbose:
                            logger.info(f"Image summary at position {task['position'] + 0.5}: {summary[:MAX_VERBOSE_LENGTH]}...")
                    task['image_bytes'] = None
                del image_tasks
            del positioned_elements, tables, doc, res
            gc.collect()
            release_memory()

            # Clear stale cached pipelines — Docling may create duplicate
            # pipelines with different option hashes due to mutable defaults
            # in OCR options being mutated during processing.
            if hasattr(converter, 'initialized_pipelines') and len(converter.initialized_pipelines) > 1:
                last_key = list(converter.initialized_pipelines.keys())[-1]
                stale_keys = [k for k in converter.initialized_pipelines if k != last_key]
                for k in stale_keys:
                    del converter.initialized_pipelines[k]
                gc.collect()
                release_memory()

        # GMFT tables (processes entire file, only once)
        if self.parse_tables and self.enable_gmft and filename.endswith('.pdf'):
            all_tables = list(self.get_tables_with_gmft(filename))

        # Load image bytes back from disk now that docling memory is freed
        all_image_bytes = []
        if img_manifest:
            for image_id, img_path in img_manifest:
                with open(img_path, 'rb') as f:
                    all_image_bytes.append((image_id, f.read()))
                os.remove(img_path)
            os.rmdir(img_tmp_dir)

        release_memory()
        return self._finalize(filename, doc_title, all_positioned, all_tables, all_image_bytes)

    def _finalize(self, filename, doc_title, positioned_elements, tables, image_bytes):
        """Sort elements and build final ParsedDocument."""
        positioned_elements.sort(key=lambda x: x[0])
        content_stream = [(content, metadata) for _, content, metadata in positioned_elements]

        if not doc_title:
            basename = os.path.basename(filename)
            doc_title = os.path.splitext(basename)[0].replace('_', ' ').replace('-', ' ').title()
            logger.info(f"No title found in document, using filename fallback: '{doc_title}' for file {filename}")

        return ParsedDocument(
            title=doc_title,
            content_stream=content_stream,
            tables=tables,
            image_bytes=image_bytes
        )

    def parse(self, filename: str, source_url: str = "No URL") -> ParsedDocument:
        """
        Parse a local file and return unified content stream with images interleaved in proper order.
        Tables are extracted separately for structured indexing.
        Uses position-based ordering to maintain document structure.

        When ``fallback_ocr`` is enabled and the initial (non-OCR) parse of a PDF
        yields no meaningful text, the file is automatically re-parsed with OCR.
        """
        st = time.time()

        converter = self._get_or_create_converter()
        result = self._parse_with_converter(filename, source_url, converter)

        # Fallback: if content stream is empty for a PDF, retry with OCR
        if (self.fallback_ocr
                and not self.do_ocr
                and filename.lower().endswith('.pdf')
                and not result.content_stream):
            logger.info(f"No content elements found in {filename}, retrying with OCR fallback")
            ocr_converter = self._get_or_create_converter(fallback_ocr=True)
            result = self._parse_with_converter(filename, source_url, ocr_converter)

        logger.info(f"DoclingParser: {len(result.content_stream)} content elements, {len(result.tables)} tables")
        logger.info(f"parsing file {filename} with Docling took {time.time()-st:.2f} seconds")

        return result



class ImageFileParser(DocumentParser):
    """
    Parser for standalone image files (PNG, JPG, GIF, etc.)
    Returns a single-element ParsedDocument with the image summary

    Note: This parser always requires summarize_images=True and vision model config.
    """
    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = None,
    ):
        model_config = model_config or {}

        # Validate that vision config key exists (value can be empty dict for testing)
        if 'vision' not in model_config:
            raise ValueError("ImageFileParser requires 'vision' key in model_config")

        super().__init__(
            cfg=cfg,
            verbose=verbose,
            model_config=model_config,
            parse_tables=False,
            enable_gmft=False,
            do_ocr=False,
            summarize_images=True  # Always True for image files
        )

    def parse(self, filename: str, source_url: str = "No URL") -> ParsedDocument:
        """
        Parse a standalone image file and return single-element ParsedDocument

        Args:
            filename: Path to the image file
            source_url: Source URL for context

        Returns:
            ParsedDocument with single image element
        """
        st = time.time()

        if not os.path.exists(filename):
            logger.error(f"Image file {filename} does not exist")
            return ParsedDocument(title='', content_stream=[], tables=[], image_bytes=[])

        # Use filename (without extension) as title
        basename = os.path.basename(filename)
        doc_title = os.path.splitext(basename)[0].replace('_', ' ').replace('-', ' ').title()

        # Use filename and path as context for standalone images
        file_context = f"Filename: {basename}\nPath: {filename}"
        if source_url != "No URL":
            file_context += f"\nSource: {source_url}"

        # Generate image summary with filename/path as context
        try:
            # Read image binary data
            with open(filename, 'rb') as fp:
                image_binary = fp.read()
            image_id = f"standalone_image_{slugify(basename)}"
            image_bytes = [(image_id, image_binary)]

            image_summary = self.image_summarizer.summarize_image(
                filename, source_url, previous_text=file_context, next_text=None
            )

            if not image_summary:
                logger.warning(f"Failed to generate summary for image {filename}")
                return ParsedDocument(title=doc_title, content_stream=[], tables=[], image_bytes=[])

            # Create single element with position 0
            metadata = {
                'element_type': 'image',
                'page': 1,
                'image_id': image_id,
                'filename': basename
            }
            content_stream = [(image_summary, metadata)]

            if self.verbose:
                logger.info(f"Image summary: {image_summary[:MAX_VERBOSE_LENGTH]}...")

            logger.info(f"ImageFileParser: parsed standalone image {filename} in {time.time()-st:.2f} seconds")

            return ParsedDocument(
                title=doc_title,
                content_stream=content_stream,
                tables=[],
                image_bytes=image_bytes
            )

        except Exception as e:
            logger.error(f"Failed to parse image file {filename}: {e}")
            return ParsedDocument(title=doc_title, content_stream=[], tables=[], image_bytes=[])


class UnstructuredDocumentParser(DocumentParser):
    def __init__(
        self,
        cfg: OmegaConf,
        verbose: bool = False,
        model_config: dict = {},
        chunking_strategy: str = "by_title",
        chunk_size: int = 1024,
        parse_tables: bool = False,
        enable_gmft: bool = False,
        summarize_images: bool = False,
        image_context: dict = None,
        hi_res_model_name: str = "yolox"
    ):
        super().__init__(
            cfg=cfg,
            verbose=verbose,
            model_config=model_config,
            parse_tables=parse_tables,
            enable_gmft=enable_gmft,
            do_ocr=False,
            summarize_images=summarize_images
        )
        self.chunking_strategy = chunking_strategy     # none, by_title or basic
        self.chunk_size = chunk_size
        self.image_context = image_context or {'num_previous_chunks': 1, 'num_next_chunks': 1}
        self.hi_res_model_name = hi_res_model_name
        if self.verbose:
            logger.info(f"Using UnstructuredDocumentParser with chunking strategy '{self.chunking_strategy}', chunk size {self.chunk_size}, hi_res_model_name '{self.hi_res_model_name}'")

    def _get_elements(
        self,
        filename: str,
        override_chunking: bool = False,
    ) -> List[us.documents.elements.Element]:
        '''
        Get elements from document using Unstructured partition_XXX functions
        Args:
            filename (str): Name of the file to parse.
            override_chunking (bool): If True, disable chunking regardless of config
        Returns:
            List of elements from the document
        '''

        mime_type = detect_file_type(filename)
        partition_kwargs = {'infer_table_structure': True}
        
        # Apply chunking strategy if configured and not overridden
        if self.chunking_strategy != 'none' and not override_chunking:
            partition_kwargs.update({
                "chunking_strategy": self.chunking_strategy,
                "max_characters": self.chunk_size,
            })
        
        # Configure image and table extraction for PDFs
        if mime_type == 'application/pdf':
            partition_kwargs.update({
                'extract_images_in_pdf': True,
                'extract_image_block_types': ["Image", "Table"],
                'strategy': "hi_res",
                'hi_res_model_name': self.hi_res_model_name,
            })

        if mime_type == 'application/pdf':
            partition_func = partition_pdf
        elif mime_type == 'text/html' or mime_type == 'application/xml':
            partition_func = partition_html
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            partition_func = partition_docx
        elif mime_type == 'application/vnd.openxmlformats-officedocument.presentationml.presentation':
            partition_func = partition_pptx
        else:
            logger.info(f"data from {filename} is not HTML, PPTX, DOCX or PDF (mime type = {mime_type}), skipping")
            return []

        try:
            elements = partition_func(
                filename=filename,
                **partition_kwargs
            )
            return elements
        except Exception as e:
            logger.error(f"Error occurred while partitioning document {filename}: {e}")
            return []

    def _get_tables(self, elements):
        table_data = []
        for e in elements:
            if isinstance(e, us.documents.elements.Table):
                try:
                    html_table = e.metadata.text_as_html
                    df = pd.read_html(StringIO(html_table))[0]
                    table_data.append((df, str(e), e.metadata.page_number))
                except Exception as err:
                    logger.error(f"Error parsing HTML table: {err}. Skipping...")
                    continue

        if not table_data:
            return []

        table_texts = [text for _, text, _ in table_data]
        summaries = self._parallel_summarize_tables(table_texts)
        return [[df, summary, '', {'page': page_num}]
                for (df, _, page_num), summary in zip(table_data, summaries)]

    def parse(self, filename: str, source_url: str = "No URL") -> ParsedDocument:
        """
        Parse a document and return unified content stream with images inline.
        Tables are extracted separately for structured indexing.
        When chunking is enabled, make two passes:
        1. Extract chunked text
        2. Extract raw images without chunking
        Then map everything together using position-based ordering.
        """
        st = time.time()

        if self.verbose:
            logger.info(f"Unstructured: extracting all content from {filename}")
        
        # Determine if we're using chunking
        is_chunking = self.chunking_strategy != 'none'
        
        if is_chunking:
            # Two-pass extraction for chunked content
            if self.verbose:
                logger.info("Using dual extraction with position tracking: chunked text + raw tables/images")
            
            # Pass 1: Get raw elements without chunking to establish positions
            raw_elements = self._get_elements(filename, override_chunking=True)
            
            # Create position map for raw elements
            raw_positions = {}
            for idx, element in enumerate(raw_elements):
                page_num = getattr(element.metadata, 'page_number', 1) or 1
                # Position formula: page_num * 1000 + element_index
                position = page_num * 1000 + idx
                raw_positions[id(element)] = (position, page_num, idx)
            
            # Pass 2: Get chunked text elements
            chunked_elements = self._get_elements(filename, override_chunking=False)
            
            # Log what we found
            chunked_types = {}
            for e in chunked_elements:
                etype = type(e).__name__
                chunked_types[etype] = chunked_types.get(etype, 0) + 1
            logger.info(f"Chunked element types: {chunked_types}")
            
            raw_types = {}
            for e in raw_elements:
                etype = type(e).__name__
                raw_types[etype] = raw_types.get(etype, 0) + 1
            logger.info(f"Raw element types: {raw_types}")
            
            elements = chunked_elements  # Use chunked for text processing
            raw_tables_images = raw_elements  # Keep raw for table/image extraction
        else:
            # Single pass when not chunking
            elements = self._get_elements(filename)
            raw_tables_images = elements  # Same elements for everything
            raw_positions = None  # Not needed when not chunking
            
            # Log element types
            element_types = {}
            for e in elements:
                etype = type(e).__name__
                element_types[etype] = element_types.get(etype, 0) + 1
            logger.info(f"Element types found: {element_types}")
        
        # Build list of positioned elements
        positioned_elements = []  # List of (position, content, metadata) tuples
        image_bytes = []  # Store image binary data
        
        if not is_chunking:
            # When not chunking, collect text inline and image tasks for parallel summarization
            image_tasks = []
            image_counter = 0
            for idx, element in enumerate(elements):
                page_num = getattr(element.metadata, 'page_number', 1) or 1
                base_position = page_num * 1000

                if isinstance(element, us.documents.elements.Image):
                    if self.summarize_images:
                        has_valid_coords = (
                            element.metadata.coordinates and
                            element.metadata.coordinates.system.width > 50 and
                            element.metadata.coordinates.system.height > 50
                        )
                        has_image_path = hasattr(element.metadata, 'image_path') and element.metadata.image_path

                        if has_image_path or has_valid_coords:
                            try:
                                previous_text, next_text = extract_image_context(
                                    elements, idx,
                                    num_previous=self.image_context['num_previous_chunks'],
                                    num_next=self.image_context['num_next_chunks']
                                )

                                image_path = getattr(element.metadata, 'image_path', None)
                                if not image_path and hasattr(element, 'image'):
                                    image_path = element.image

                                if image_path:
                                    image_binary = None
                                    image_id = None
                                    if os.path.exists(image_path):
                                        if self.store_image_bytes:
                                            with open(image_path, 'rb') as fp:
                                                image_binary = fp.read()
                                        image_id = f"unstructured_page_{page_num}_image_{image_counter}"
                                        image_counter += 1

                                    image_tasks.append({
                                        'image_path': image_path, 'source_url': source_url,
                                        'previous_text': previous_text, 'next_text': next_text,
                                        'image_bytes': image_binary,
                                        'image_id': image_id,
                                        'position': base_position + idx + 0.5,
                                        'page_no': page_num,
                                    })
                                else:
                                    logger.warning("Image element found but no valid image path available")
                            except Exception as exc:
                                logger.error(f"Error collecting image data: {exc}")

                elif isinstance(element, us.documents.elements.Table):
                    pass
                else:
                    position = base_position + idx
                    metadata = {'element_type': 'text', 'page': page_num}
                    positioned_elements.append((position, str(element), metadata))

            # Parallel image summarization (non-chunking path)
            if image_tasks:
                summaries = self._parallel_summarize_images(image_tasks)
                for task, summary in zip(image_tasks, summaries):
                    if self.store_image_bytes and task['image_bytes'] and task['image_id']:
                        image_bytes.append((task['image_id'], task['image_bytes']))
                    if summary:
                        metadata = {
                            'element_type': 'image',
                            'page': task['page_no'],
                            'image_id': task['image_id']
                        }
                        positioned_elements.append((task['position'], summary, metadata))
                        if self.verbose:
                            logger.info(f"Image summary at position {task['position']}: {summary[:MAX_VERBOSE_LENGTH]}...")
                    task['image_bytes'] = None  # Release memory
        else:
            # When chunking, process chunked text with position mapping
            chunk_position_counter = {}
            for element in elements:
                page_num = getattr(element.metadata, 'page_number', 1) or 1
                
                # Skip tables and images in chunked elements (will process from raw)
                if isinstance(element, (us.documents.elements.Table, us.documents.elements.Image)):
                    continue
                
                # Process text/composite elements
                # Assign positions based on page and order within page
                if page_num not in chunk_position_counter:
                    chunk_position_counter[page_num] = 0
                
                base_position = page_num * 1000
                position = base_position + chunk_position_counter[page_num]
                chunk_position_counter[page_num] += 10  # Leave gaps for inserting images/tables
                
                metadata = {'element_type': 'text', 'page': page_num}
                positioned_elements.append((position, str(element), metadata))
        
            # Collect image tasks from raw extraction for parallel summarization
            image_tasks = []
            image_counter = 0
            for idx, element in enumerate(raw_tables_images):
                page_num = getattr(element.metadata, 'page_number', 1) or 1
                base_position = page_num * 1000

                if isinstance(element, us.documents.elements.Image):
                    logger.info(f"Found Image element on page {page_num}")
                    logger.info(f"  Has coordinates: {hasattr(element.metadata, 'coordinates')}")
                    if hasattr(element.metadata, 'coordinates') and element.metadata.coordinates:
                        logger.info(f"  Coords dimensions: {element.metadata.coordinates.system.width}x{element.metadata.coordinates.system.height}")
                    logger.info(f"  Has image_path: {hasattr(element.metadata, 'image_path')}")
                    if hasattr(element.metadata, 'image_path'):
                        logger.info(f"  Image path: {element.metadata.image_path}")

                    if self.summarize_images:
                        has_valid_coords = (
                            element.metadata.coordinates and
                            element.metadata.coordinates.system.width > 50 and
                            element.metadata.coordinates.system.height > 50
                        )
                        has_image_path = hasattr(element.metadata, 'image_path') and element.metadata.image_path

                        if has_image_path or has_valid_coords:
                            try:
                                image_position = base_position + idx + 0.5

                                best_context_idx = 0
                                for chunk_idx, chunk_elem in enumerate(elements):
                                    chunk_page = getattr(chunk_elem.metadata, 'page_number', 1) or 1
                                    chunk_position = chunk_page * 1000 + chunk_idx
                                    if chunk_position > image_position:
                                        best_context_idx = chunk_idx
                                        break
                                    best_context_idx = chunk_idx

                                previous_text, next_text = extract_image_context(
                                    elements, best_context_idx,
                                    num_previous=self.image_context['num_previous_chunks'],
                                    num_next=self.image_context['num_next_chunks']
                                )

                                image_path = getattr(element.metadata, 'image_path', None)
                                if not image_path and hasattr(element, 'image'):
                                    image_path = element.image

                                if image_path:
                                    image_binary = None
                                    image_id = None
                                    if os.path.exists(image_path):
                                        if self.store_image_bytes:
                                            with open(image_path, 'rb') as fp:
                                                image_binary = fp.read()
                                        image_id = f"unstructured_chunked_page_{page_num}_image_{image_counter}"
                                        image_counter += 1

                                    image_tasks.append({
                                        'image_path': image_path, 'source_url': source_url,
                                        'previous_text': previous_text, 'next_text': next_text,
                                        'image_bytes': image_binary,
                                        'image_id': image_id,
                                        'position': image_position,
                                        'page_no': page_num,
                                    })
                                else:
                                    logger.warning("Image element found but no valid image path available")
                            except Exception as exc:
                                logger.error(f"Error collecting image data: {exc}")
                                continue

                elif isinstance(element, us.documents.elements.Table):
                    if self.verbose:
                        logger.info(f"Table found on page {page_num} - will be processed for structured indexing")

            # Parallel image summarization (chunking path)
            if image_tasks:
                summaries = self._parallel_summarize_images(image_tasks)
                for task, summary in zip(image_tasks, summaries):
                    if self.store_image_bytes and task['image_bytes'] and task['image_id']:
                        image_bytes.append((task['image_id'], task['image_bytes']))
                    if summary:
                        metadata = {
                            'element_type': 'image',
                            'page': task['page_no'],
                            'image_id': task['image_id']
                        }
                        positioned_elements.append((task['position'], summary, metadata))
                        if self.verbose:
                            logger.info(f"Image summary at position {task['position']}: {summary[:MAX_VERBOSE_LENGTH]}...")
                    task['image_bytes'] = None  # Release memory

        # Sort all positioned elements by position to create final document order
        positioned_elements.sort(key=lambda x: x[0])
        
        # Convert to content_stream format (without positions)
        content_stream = [(content, metadata) for _, content, metadata in positioned_elements]
        
        if self.verbose:
            logger.info(f"Final content stream has {len(content_stream)} elements")
            # Log the element type distribution in final stream
            type_counts = {}
            for _, metadata in content_stream:
                etype = metadata.get('element_type', 'unknown')
                type_counts[etype] = type_counts.get(etype, 0) + 1
            logger.info(f"Final element distribution: {type_counts}")
        
        # Find document title with priority: Document metadata > Title elements > filename
        doc_title = extract_document_title(filename)
        
        # For files without metadata title, look for Title elements (except PDF/DOCX/PPTX which skip Title elements)
        if not doc_title and not filename.endswith(('.pdf', '.docx', '.pptx')):
            title_elements = raw_tables_images if is_chunking else elements
            titles = [str(x) for x in title_elements if type(x) == us.documents.elements.Title and len(str(x).strip()) > 3]
            if titles:
                doc_title = titles[0]
                logger.info(f"Extracted document title from Title element: '{doc_title}' from file {filename}")
        
        # Fallback: use filename if no title found
        if not doc_title:
            basename = os.path.basename(filename)
            doc_title = os.path.splitext(basename)[0].replace('_', ' ').replace('-', ' ').title()
            logger.info(f"No title found in document, using filename fallback: '{doc_title}' for file {filename}")

        # Process tables separately for structured indexing
        tables = []
        if self.parse_tables:
            if self.enable_gmft and filename.endswith('.pdf'):
                tables = list(self.get_tables_with_gmft(filename))
            else:
                # Use raw elements for table extraction when chunking
                table_elements = raw_tables_images if is_chunking else elements
                tables = list(self._get_tables(table_elements))

        # Release heavy element lists now that extraction is complete
        del raw_tables_images, elements
        if is_chunking:
            del raw_elements, chunked_elements
        gc.collect()

        logger.info(f"UnstructuredParser: {len(content_stream)} content elements, {len(tables)} tables")
        logger.info(f"parsing file {filename} with unstructured.io took {time.time()-st:.2f} seconds")

        return ParsedDocument(
            title=doc_title,
            content_stream=content_stream,
            tables=tables,
            image_bytes=image_bytes
        )

